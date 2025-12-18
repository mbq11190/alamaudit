# -*- coding: utf-8 -*-

import io
import logging
import os
from datetime import datetime

from odoo import http
from odoo.http import request, content_disposition

_logger = logging.getLogger(__name__)

# Try to import optional PDF generation libraries
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    _logger.info("WeasyPrint not available - PDF generation will use alternative method")

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    _logger.info("python-docx not available - Word generation will use alternative method")


class FitProperTemplateController(http.Controller):
    """Controller for downloading Fit & Proper Assessment templates with auto-fill."""

    def _get_template_path(self):
        """Get the path to the HTML template."""
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(module_path, 'static', 'templates', 'fit_proper_assessment_template.html')

    def _get_onboarding_data(self, onboarding_id):
        """Get client name and period from onboarding record."""
        data = {
            'entity_name': '',
            'audit_period': '',
            'reference': '',
            'generated_date': datetime.now().strftime('%d %B %Y'),
        }
        
        if onboarding_id:
            try:
                onboarding = request.env['qaco.client.onboarding'].sudo().browse(int(onboarding_id))
                if onboarding.exists():
                    # Entity/Client Name
                    data['entity_name'] = onboarding.client_id.name or ''
                    
                    # Audit Period from audit_year
                    if onboarding.audit_id and onboarding.audit_id.audit_year:
                        years = onboarding.audit_id.audit_year.mapped('name')
                        data['audit_period'] = ', '.join(years) if years else ''
                    
                    # Reference
                    data['reference'] = f"FP-{onboarding.audit_id.name or 'N/A'}" if onboarding.audit_id else ''
            except Exception as e:
                _logger.warning(f"Could not fetch onboarding data: {e}")
        
        return data

    def _read_html_template(self, data=None):
        """Read and populate the HTML template file."""
        template_path = self._get_template_path()
        with open(template_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Replace placeholders with actual data
        if data:
            html_content = html_content.replace('{{ENTITY_NAME}}', data.get('entity_name', '') or '_________________________')
            html_content = html_content.replace('{{AUDIT_PERIOD}}', data.get('audit_period', '') or '_________________________')
            html_content = html_content.replace('{{REFERENCE}}', data.get('reference', '') or '_________________________')
            html_content = html_content.replace('{{GENERATED_DATE}}', data.get('generated_date', datetime.now().strftime('%d %B %Y')))
        else:
            # Blank template
            html_content = html_content.replace('{{ENTITY_NAME}}', '_________________________')
            html_content = html_content.replace('{{AUDIT_PERIOD}}', '_________________________')
            html_content = html_content.replace('{{REFERENCE}}', '_________________________')
            html_content = html_content.replace('{{GENERATED_DATE}}', datetime.now().strftime('%d %B %Y'))
        
        return html_content

    @http.route('/onboarding/template/fit_proper/html', type='http', auth='user', methods=['GET'])
    def download_fit_proper_html(self, onboarding_id=None, **kwargs):
        """Download Fit & Proper template as HTML with optional auto-fill."""
        try:
            data = self._get_onboarding_data(onboarding_id)
            html_content = self._read_html_template(data)
            
            # Create filename with client name if available
            filename = 'Fit_Proper_Assessment'
            if data.get('entity_name'):
                safe_name = "".join(c for c in data['entity_name'] if c.isalnum() or c in (' ', '-', '_')).strip()[:30]
                filename = f"Fit_Proper_{safe_name.replace(' ', '_')}"
            
            return request.make_response(
                html_content,
                headers=[
                    ('Content-Type', 'text/html; charset=utf-8'),
                    ('Content-Disposition', content_disposition(f'{filename}.html')),
                ]
            )
        except Exception as e:
            _logger.error(f"Error generating HTML template: {e}")
            return request.not_found()

    @http.route('/onboarding/template/fit_proper/pdf', type='http', auth='user', methods=['GET'])
    def download_fit_proper_pdf(self, onboarding_id=None, **kwargs):
        """Download Fit & Proper template as PDF with optional auto-fill."""
        try:
            data = self._get_onboarding_data(onboarding_id)
            html_content = self._read_html_template(data)
            
            # Create filename with client name if available
            filename = 'Fit_Proper_Assessment'
            if data.get('entity_name'):
                safe_name = "".join(c for c in data['entity_name'] if c.isalnum() or c in (' ', '-', '_')).strip()[:30]
                filename = f"Fit_Proper_{safe_name.replace(' ', '_')}"
            
            if WEASYPRINT_AVAILABLE:
                # Use WeasyPrint for high-quality PDF
                pdf_buffer = io.BytesIO()
                HTML(string=html_content).write_pdf(pdf_buffer)
                pdf_content = pdf_buffer.getvalue()
            else:
                # Fallback: Use Odoo's built-in wkhtmltopdf if available
                try:
                    from odoo.tools import pdf
                    pdf_content = pdf.html2pdf(html_content)
                except Exception:
                    # Last resort: Return HTML with PDF content type (browser will render)
                    _logger.warning("No PDF generator available, returning HTML for print-to-PDF")
                    return request.make_response(
                        html_content,
                        headers=[
                            ('Content-Type', 'text/html; charset=utf-8'),
                            ('Content-Disposition', content_disposition(f'{filename}_PrintToPDF.html')),
                        ]
                    )
            
            return request.make_response(
                pdf_content,
                headers=[
                    ('Content-Type', 'application/pdf'),
                    ('Content-Disposition', content_disposition(f'{filename}.pdf')),
                ]
            )
        except Exception as e:
            _logger.error(f"Error generating PDF template: {e}")
            return request.not_found()

    @http.route('/onboarding/template/fit_proper/docx', type='http', auth='user', methods=['GET'])
    def download_fit_proper_docx(self, onboarding_id=None, **kwargs):
        """Download Fit & Proper template as Word document with optional auto-fill."""
        try:
            data = self._get_onboarding_data(onboarding_id)
            
            # Create filename with client name if available
            filename = 'Fit_Proper_Assessment'
            if data.get('entity_name'):
                safe_name = "".join(c for c in data['entity_name'] if c.isalnum() or c in (' ', '-', '_')).strip()[:30]
                filename = f"Fit_Proper_{safe_name.replace(' ', '_')}"
            
            if DOCX_AVAILABLE:
                doc = self._create_word_document(data)
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_content = docx_buffer.getvalue()
            else:
                # Fallback: Return HTML with docx-friendly styling
                _logger.warning("python-docx not available, returning HTML for Word import")
                html_content = self._read_html_template(data)
                return request.make_response(
                    html_content,
                    headers=[
                        ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                        ('Content-Disposition', content_disposition(f'{filename}.doc')),
                    ]
                )
            
            return request.make_response(
                docx_content,
                headers=[
                    ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    ('Content-Disposition', content_disposition(f'{filename}.docx')),
                ]
            )
        except Exception as e:
            _logger.error(f"Error generating Word template: {e}")
            return request.not_found()

    def _create_word_document(self, data=None):
        """Create a Word document for the Fit & Proper template with auto-fill."""
        doc = Document()
        
        # Set document margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Title
        title = doc.add_heading('📄 FIT & PROPER ASSESSMENT & CONFIRMATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('(Consolidated Master Template)')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True
        
        # Header Info Box
        entity_name = data.get('entity_name', '') if data else ''
        audit_period = data.get('audit_period', '') if data else ''
        reference = data.get('reference', '') if data else ''
        generated_date = data.get('generated_date', datetime.now().strftime('%d %B %Y')) if data else datetime.now().strftime('%d %B %Y')
        
        header_table = doc.add_table(rows=2, cols=4)
        header_table.style = 'Table Grid'
        header_table.rows[0].cells[0].text = 'Entity Name:'
        header_table.rows[0].cells[1].text = entity_name or '_________________________'
        header_table.rows[0].cells[2].text = 'Audit Period:'
        header_table.rows[0].cells[3].text = audit_period or '_________________________'
        header_table.rows[1].cells[0].text = 'Generated:'
        header_table.rows[1].cells[1].text = generated_date
        header_table.rows[1].cells[2].text = 'Reference:'
        header_table.rows[1].cells[3].text = reference or '_________________________'
        doc.add_paragraph()
        
        # Section A: Entity & Individual Identification
        doc.add_heading('A. ENTITY & INDIVIDUAL IDENTIFICATION', level=1)
        table_a = doc.add_table(rows=10, cols=2)
        table_a.style = 'Table Grid'
        
        id_items = [
            ('Entity Name', entity_name or '_________________________'),
            ('Legal Status', '☐ Private  ☐ Listed  ☐ NGO / NPO  ☐ Other'),
            ('Registration / NTN', '_________________________'),
            ('Individual Name', '_________________________'),
            ('CNIC / Passport No.', '_________________________'),
            ('Nationality', '_________________________'),
            ('Role Assessed', '☐ Director  ☐ CEO  ☐ CFO  ☐ Key Management  ☐ Trustee'),
            ('Date of Appointment', '_________________________'),
            ('Date of Assessment', '_________________________'),
            ('Applicable Regulator', '☐ SECP  ☐ SBP  ☐ NGO  ☐ Other: ___________'),
        ]
        for i, (item, value) in enumerate(id_items):
            table_a.rows[i].cells[0].text = item
            table_a.rows[i].cells[1].text = value
        
        # Section B: Integrity, Honesty & Reputation
        doc.add_heading('B. INTEGRITY, HONESTY & REPUTATION ASSESSMENT', level=1)
        p = doc.add_paragraph('(Mandatory – Companies Act 2017 / ISQM-1 / IESBA)')
        p.runs[0].italic = True
        
        table_b = doc.add_table(rows=6, cols=4)
        table_b.style = 'Table Grid'
        table_b.rows[0].cells[0].text = 'Criteria'
        table_b.rows[0].cells[1].text = 'Yes'
        table_b.rows[0].cells[2].text = 'No'
        table_b.rows[0].cells[3].text = 'Remarks'
        
        integrity_items = [
            'No criminal conviction involving fraud, dishonesty, or moral turpitude',
            'No pending prosecution or investigation',
            'No adverse regulatory enforcement or penalty',
            'Not disqualified under Companies Act, 2017',
            'No history of willful default or financial misconduct',
        ]
        for i, item in enumerate(integrity_items, 1):
            table_b.rows[i].cells[0].text = item
            table_b.rows[i].cells[1].text = '☐'
            table_b.rows[i].cells[2].text = '☐'
        
        # Section C: Competence & Experience
        doc.add_heading('C. COMPETENCE & EXPERIENCE ASSESSMENT', level=1)
        
        table_c = doc.add_table(rows=6, cols=4)
        table_c.style = 'Table Grid'
        table_c.rows[0].cells[0].text = 'Criteria'
        table_c.rows[0].cells[1].text = 'Yes'
        table_c.rows[0].cells[2].text = 'No'
        table_c.rows[0].cells[3].text = 'Remarks'
        
        competence_items = [
            'Relevant academic qualification',
            'Relevant professional / industry experience',
            "Adequate understanding of entity's business",
            'Knowledge of applicable laws & regulations',
            'Financial literacy (mandatory for directors / audit committee)',
        ]
        for i, item in enumerate(competence_items, 1):
            table_c.rows[i].cells[0].text = item
            table_c.rows[i].cells[1].text = '☐'
            table_c.rows[i].cells[2].text = '☐'
        
        # Section D: Financial Soundness
        doc.add_heading('D. FINANCIAL SOUNDNESS', level=1)
        
        table_d = doc.add_table(rows=5, cols=4)
        table_d.style = 'Table Grid'
        table_d.rows[0].cells[0].text = 'Criteria'
        table_d.rows[0].cells[1].text = 'Yes'
        table_d.rows[0].cells[2].text = 'No'
        table_d.rows[0].cells[3].text = 'Remarks'
        
        financial_items = [
            'Not declared insolvent or bankrupt',
            'No overdue taxes or statutory defaults',
            'No material unpaid liabilities to lenders',
            'No adverse credit information',
        ]
        for i, item in enumerate(financial_items, 1):
            table_d.rows[i].cells[0].text = item
            table_d.rows[i].cells[1].text = '☐'
            table_d.rows[i].cells[2].text = '☐'
        
        # Section E: Independence & Conflict
        doc.add_heading('E. INDEPENDENCE & CONFLICT OF INTEREST', level=1)
        
        table_e = doc.add_table(rows=5, cols=4)
        table_e.style = 'Table Grid'
        table_e.rows[0].cells[0].text = 'Criteria'
        table_e.rows[0].cells[1].text = 'Yes'
        table_e.rows[0].cells[2].text = 'No'
        table_e.rows[0].cells[3].text = 'Remarks'
        
        independence_items = [
            "No conflict with entity's interests",
            'No prohibited related-party relationships',
            'All actual / potential conflicts disclosed',
            'Not holding incompatible positions',
        ]
        for i, item in enumerate(independence_items, 1):
            table_e.rows[i].cells[0].text = item
            table_e.rows[i].cells[1].text = '☐'
            table_e.rows[i].cells[2].text = '☐'
        
        # Section F: AML/CFT, PEP & Sanctions
        doc.add_heading('F. AML / CFT, PEP & SANCTIONS SCREENING', level=1)
        p = doc.add_paragraph('(Mandatory – AML Act / ISQM-1)')
        p.runs[0].italic = True
        
        table_f = doc.add_table(rows=5, cols=3)
        table_f.style = 'Table Grid'
        table_f.rows[0].cells[0].text = 'Check'
        table_f.rows[0].cells[1].text = 'Status'
        table_f.rows[0].cells[2].text = 'Evidence Reference'
        
        table_f.rows[1].cells[0].text = 'Politically Exposed Person (PEP)'
        table_f.rows[1].cells[1].text = '☐ Clear  ☐ Identified'
        table_f.rows[1].cells[2].text = 'Screening report'
        
        table_f.rows[2].cells[0].text = 'Sanctions (UN / OFAC / EU)'
        table_f.rows[2].cells[1].text = '☐ Clear  ☐ Flagged'
        table_f.rows[2].cells[2].text = 'Screening report'
        
        table_f.rows[3].cells[0].text = 'Country Risk'
        table_f.rows[3].cells[1].text = '☐ Low  ☐ Medium  ☐ High'
        table_f.rows[3].cells[2].text = 'Narrative'
        
        table_f.rows[4].cells[0].text = 'Source of Wealth / Income'
        table_f.rows[4].cells[1].text = '☐ Reasonable  ☐ Concern'
        table_f.rows[4].cells[2].text = 'Explanation'
        
        # Section G: EDD
        doc.add_heading('G. ENHANCED DUE DILIGENCE (EDD) – IF APPLICABLE', level=1)
        doc.add_paragraph('EDD Trigger(s):').bold = True
        doc.add_paragraph('☐ PEP    ☐ Foreign national (high-risk jurisdiction)    ☐ Complex ownership / nominee structure    ☐ Adverse media')
        doc.add_paragraph()
        doc.add_paragraph('EDD Procedures Performed:').bold = True
        doc.add_paragraph('_' * 80)
        doc.add_paragraph('_' * 80)
        
        # Section H: Overall Assessment
        doc.add_heading('H. OVERALL FIT & PROPER ASSESSMENT (AUDITOR / FIRM)', level=1)
        
        doc.add_paragraph('Assessment Outcome:').bold = True
        doc.add_paragraph('☐ Fit & Proper')
        doc.add_paragraph('☐ Fit & Proper with Conditions')
        doc.add_paragraph('☐ Not Fit & Proper')
        doc.add_paragraph()
        doc.add_paragraph('Conditions / Safeguards (if any):').bold = True
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        doc.add_paragraph('Conclusion:').bold = True
        doc.add_paragraph('Based on the above assessment, the individual meets / does not meet the Fit & Proper criteria under applicable laws, regulations, and ethical requirements.')
        
        # Section I: Declaration
        doc.add_heading('I. DECLARATION BY THE INDIVIDUAL (CONFIRMATION)', level=1)
        
        doc.add_paragraph('I hereby confirm that:').bold = True
        doc.add_paragraph('• The information provided above is true, complete, and accurate.')
        doc.add_paragraph('• I have disclosed all relevant information, including conflicts of interest and regulatory matters.')
        doc.add_paragraph('• I undertake to immediately inform the Company and auditors of any material change affecting my Fit & Proper status.')
        doc.add_paragraph()
        doc.add_paragraph('Name: ________________________________     Date: ________________')
        doc.add_paragraph('Signature: ________________________________')
        
        # Section J: Auditor Confirmation
        doc.add_heading('J. AUDITOR / FIRM CONFIRMATION & SIGN-OFF', level=1)
        
        doc.add_paragraph('We confirm that the Fit & Proper assessment and confirmation has been conducted and reviewed in accordance with:').bold = True
        doc.add_paragraph('• Companies Act, 2017')
        doc.add_paragraph('• ISQM-1')
        doc.add_paragraph('• IESBA Code of Ethics')
        doc.add_paragraph('• SECP / SBP / applicable regulatory requirements')
        doc.add_paragraph()
        doc.add_paragraph('Engagement Partner: ________________________________     Date: ________________')
        doc.add_paragraph('Signature: ________________________________')
        
        return doc
